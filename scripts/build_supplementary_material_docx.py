"""Build the Supplementary Material Word document from supplementary Excel files."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SUPPLEMENTARY_DIR = PROJECT_ROOT / "supplementary"
OUTPUT_FILE = SUPPLEMENTARY_DIR / "Supplementary_Material.docx"


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.bold = bold


def add_table(document: Document, dataframe: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, column in enumerate(dataframe.columns):
        table.rows[0].cells[index].text = str(column)
    for _, row in dataframe.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(dataframe.columns):
            cells[index].text = str(row[column])


def main() -> None:
    SUPPLEMENTARY_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    add_paragraph(document, "Supplementary Material", bold=True)
    add_paragraph(
        document,
        "Artificial intelligence-assisted prioritization of nanoparticle formulations for efficient tumor delivery",
        bold=True,
    )
    add_paragraph(document, "Section 1. Supplementary Tables", bold=True)
    add_paragraph(document, "Table S1. Summary of supplementary files.", bold=True)
    add_table(
        document,
        pd.DataFrame(
            [
                ["Supplementary Excel File A", "Dataset, split labels, endpoint definition, and data dictionary."],
                ["Supplementary Excel File B", "Model evaluation, optimized hyperparameters, predictions, and reports."],
                ["Supplementary Excel File C", "Virtual screening candidates and local working-range outputs."],
            ],
            columns=["File", "Content"],
        ),
    )
    add_paragraph(document, "Section 2. Supplementary Figures", bold=True)
    add_paragraph(
        document,
        "No additional supplementary figures are provided. All primary visual results are presented in the main figures.",
    )
    add_paragraph(document, "Section 3. Supplementary Data and Code Availability", bold=True)
    add_paragraph(
        document,
        "The source code will be made available in a public GitHub repository after repository finalization.",
    )
    document.save(OUTPUT_FILE)
    print(f"Generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
